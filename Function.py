import sys, os, re, openpyxl, json, requests, glob
from docx import Document
from docx.shared import Pt
from win32com import client
from openpyxl import Workbook as wb
from PyPDF2 import PdfFileMerger, PdfFileWriter, PdfFileReader
from bs4 import BeautifulSoup, NavigableString, Tag
from datetime import datetime


def test():


	projekt_ID = '21/139'
	project = projekt(projekt_ID)
	info = project.search_project()
	
	# ZAD = zadani(info)
	# zad_info = ZAD.info_zadani()
	# info.update(zad_info)
	print(info)
	# doc = doc_manipulation(info)
	# a = doc.pruvodni_zprava()

	
	

	return

class projekt:
	def __init__(self, cislo_projektu):
		self.cislo_projektu = cislo_projektu
		with open("setting.json","r",encoding = 'utf-8') as f:
			file = f.read()
			self.setting = json.loads(file)

	
		

	###################### Hledani nazvu search_project() ############################	
	def search_project(self):

		informace_o_projektu = {}

		#input project number and split
		code_project = self.cislo_projektu # input() cislo projektu 
		name_project_list = [n for n in map(int, code_project.split('/'))]
		
		projekt_rok = name_project_list[0]
		projekt_cislo = name_project_list[1]

		user_name = os.getlogin() #get user name
		
		file = os.getcwd()
		above1 = os.path.dirname(file)
		above2 = os.path.dirname(above1)

		
		#File directories	
		 
		if 'dokumenty' in above2 or 'Dokumenty' in above2 and user_name in above2 and 'R-built' in above2:
			project_directory = above2
			
		else:
			project_directory = self.setting["projekt"]["project_directory"].replace("{user_name}", user_name)
			cloud_path = False

			
		
		seznam_projektu_excel = project_directory + self.setting["projekt"]["seznam_projektu_excel"]
		Titulni_strany_file = project_directory + self.setting["projekt"]["Titulni_strany_file"]

		project_directory = project_directory + rf'\{self.setting["rok_nazvy"][f"{projekt_rok}"]}'
		
		if not os.path.exists(project_directory):
			return 'Directory does not exists'
		

		#get name project in list
		for index, filename in enumerate(os.listdir(project_directory)):
			try:
				if int(filename[0:3]) == projekt_cislo:			#project num, name & code
					list_jmeno_projektu = filename.split('_')

					#  cislo_projektu
					informace_o_projektu["cislo_projektu"] = list_jmeno_projektu[0]

					# kod_projektu
					informace_o_projektu["kod_projektu"] = list_jmeno_projektu[1]

					#rok projektu
					informace_o_projektu["rok_projektu"] = self.setting["rok_nazvy"][f"{projekt_rok}"].split()[1]

					project_f_dir = project_directory + rf'\{filename}'

					informace_o_projektu["project_f_dir"] = project_f_dir

			except: #skip errors in name  
				pass


		wb = openpyxl.load_workbook(seznam_projektu_excel, read_only=True) #load seznam projectu file
		ws = wb.sheetnames

		for index,sheet in enumerate(ws): #find the rigth sheet and set it as active
			if f"M_20{projekt_rok}" == sheet:
				wb.active = index
				ws= wb.active
			
		####  Najit informaci o projektu v 'Seznam_projektu' ####
		for row in ws.rows:	
			for cell in row:
				if cell.value == informace_o_projektu.get("kod_projektu"):
					# global jmeno_projektu

					jmeno_projektu = ws.cell(row=cell.row, column=3).value		
					jmeno_projektanta = ws.cell(row=cell.row, column=10).value
					

					if len(jmeno_projektu)>0:
						informace_o_projektu["nazev_projektu"] = jmeno_projektu
						# print(jmeno_projektu)
					else:
						print("Nenalezeno nazev projektu")

					if len(jmeno_projektanta)>0:
						informace_o_projektu["jmeno_projektanta"] = jmeno_projektanta

					else:
						print("Nenalezeno jmeno prokejektanta")
		wb.close()
		##################################################

		informace_o_projektu["jmeno_projektanta"] = self.setting["jmeno_projektanta"][informace_o_projektu.get("jmeno_projektanta")]
		informace_o_projektu["Tel_cislo"] = self.setting["Tel_cislo"][informace_o_projektu["jmeno_projektanta"]]



		############### Projektant informace ###############
		# for filename in os.listdir(Titulni_strany_file):
		# 	if informace_o_projektu.get("jmeno_projektanta") in filename:
		# 		wb = openpyxl.load_workbook(Titulni_strany_file+rf"\{filename}\Titulní strany.xlsx", read_only=True)
		# 		ws = wb.sheetnames

		# 		for index,sheet in enumerate(ws): #find the rigth sheet and set it as active
		# 			if "ZADÁNÍ" == sheet:
		# 				wb.active = index
		# 				ws= wb.active

		# 		for row in ws.rows:	
		# 			for cell in row:
		# 				if cell.value == "Projektant":
		# 					informace_o_projektu["jmeno_projektanta"] = ws.cell(row=cell.row, column=3).value
		# 				elif cell.value == "Tel. číslo":
		# 					tel = str(ws.cell(row=cell.row, column=3).value)
		# 					informace_o_projektu["Tel_cislo"] = " ".join([tel[0:3],tel[3:6],tel[6:9]])
		self.informace_o_projektu = informace_o_projektu
		return self.informace_o_projektu
	########################################################

	def get_k_u_ID(self,informace_o_projektu):
		informace_jmena = informace_o_projektu.get("nazev_projektu")	#parsel nazev mesta
		informace_jmena = informace_jmena.replace(" - ","-")	#odstraneni mezery
		informace_jmena = re.split('[-_,]',informace_jmena)		#split 
		print(informace_jmena, informace_jmena[1])

		if len(informace_jmena[0]) > 2:
			informace_jmena.insert(0," ")
			nazev_mesta = informace_jmena[1]
		else:
			nazev_mesta = informace_jmena[1]

		nazev_mesta = nazev_mesta.strip()

		if "Ústí n.L" in nazev_mesta:
			nazev_mesta = nazev_mesta.replace("Ústí n.L.","Ústí nad Labem")

		with open('k_u_ID.json','r') as f:
			file = json.load(f)

		try:	
			ID_k_u = file[f'{nazev_mesta}']
			return ID_k_u, nazev_mesta
		except:
			print ("katastrální území nenalezen")
			return None

	def hledani_info_mesta(self):
		informace_o_projektu = self.search_project()

		if informace_o_projektu =='Directory does not exists':
			return 'Directory does not exists'

		ID_k_u = self.get_k_u_ID(informace_o_projektu)

		info_webscrappe = self.webscrappe(ID_k_u, informace_o_projektu)

		if info_webscrappe == None or ID_k_u[1] not in info_webscrappe['ku']:
			info_webscrappe = self.webscrappe_backup(ID_k_u)

		informace_o_projektu.update(info_webscrappe)

		return informace_o_projektu


	################## Hledani info mesta z ceske katastralni urad ##############################
	def webscrappe(self, ID_k_u, informace_o_projektu):

		# webscrapping
		URL = self.setting["projekt"]["URL1"].replace("{ID_k_u}", str(ID_k_u[0]))	#web link bez ID k.u. 

		page = requests.get(URL)										#web scrape informaci
		soup = BeautifulSoup(page.content, "html.parser")
		results = soup.find_all("td", limit = 14)

		dict_o_mestu = {}
		dict_o_mestu.setdefault('k.u.', ID_k_u[1])

		for index, element in enumerate(results):						#loop element pro content
			if str(element.contents[0]) in ('Kraj', 'Okres', 'Obec'):
				dict_o_mestu.setdefault(element.contents[0], str(results[index+1].contents[0].split("-")[1]))

				# print(element.contents[0], str(results[index+1].contents[0].split("-")[1]))
		
		if dict_o_mestu['Kraj'] in ['Ústecký', 'Středočeský']:
			print ("Nenalezen katastrální uzemí")
			return informace_o_projektu
		else:	
			informace_o_projektu.update(dict_o_mestu)
			informace_o_projektu['misto_stavby'] = f"Obec{dict_o_mestu['Obec']}, okres{dict_o_mestu['Okres']},{dict_o_mestu['Kraj']} kraj"
			informace_o_projektu['ku'] = f"k.ú. {dict_o_mestu['k.u.']}, okr. {dict_o_mestu['Okres']}"
			return informace_o_projektu

		informace_o_projektu.update(dict_o_mestu)
		return informace_o_projektu

	######################## Hledani info mesta z jineho zdroje ###################
	def webscrappe_backup(self, ID_k_u, informace_o_projektu):

		URL = self.setting["projekt"]["URL2"].replace("{ID_k_u}", str(ID_k_u))
	
		page = requests.get(URL)										#web scrape informaci
		soup = BeautifulSoup(page.content, "html.parser")
		results = soup.find("table", class_ = 'pd padall rowcl l colwidth')

		dict_o_mestu ={}
		dict_o_mestu.setdefault('k.u.', nazev_mesta)
		for tr in results.find_all("tr",limit = 9):
			try:										#pokud se najde key a dontent v tr 
				key = tr.find("td").text				#tak vlozeni do dict
				content = tr.find('a').contents[0].text		
				dict_o_mestu.setdefault(key, content)
			except:
				pass

		if dict_o_mestu['Kraj'] not in ['Ústecký', 'Středočeský']:
			print ("Nenalezen katastrální uzemí")
			return informace_o_projektu
		else:
			informace_o_projektu.update(dict_o_mestu)
			informace_o_projektu['misto_stavby'] = f"Obec {dict_o_mestu['Obec']}, okres {dict_o_mestu['Okres']}, {dict_o_mestu['Kraj']} kraj"
			informace_o_projektu['ku'] = f"k.ú. {dict_o_mestu['k.u.']}, okr. {dict_o_mestu['Okres']}"
				
		return informace_o_projektu
		
#############################################################################



###############################################################################################################################################################



######################################

class doc_manipulation:

	def __init__(self,info_dict):
		self.project_dict = info_dict
		self.user_name = os.getlogin()
		with open("setting.json","r",encoding = 'utf-8') as f:
			file = f.read()
			self.setting = json.loads(file)
	

	def doc_replace_fuc(self,dic_new,dic_old,key,place):
		if str(dic_old.get(f"{key}")) in place.text:
			inline = place.runs

			for i in range(len(inline)):							#text
				if  str(dic_old.get(f"{key}")) in inline[i].text :
					# print(f'Found: {inline[i].text}',  inline[i].text in str(dic_old.get(f"{key}")) )
					# print('Replace from:' ,inline[i].text,)
					inline[i].text = inline[i].text.replace(str(dic_old.get(f"{key}")),str(dic_new.get(f"{key}")))
					# print('To: ',inline[i].text + '\n')
			
		
	def pruvodni_zprava(self):

		user_name = self.user_name
	
		cwd = os.getcwd()
		
		
		today = datetime.now()
		today = today.strftime("%m %Y")
		today = today.replace(" ","/")

		self.project_dict["datum_vydani"] = today
		stary_veci = self.setting["pruvodni_zprava"]
		
		def sub_pruvodni_zprava(word_file):

			try:
				with open(fr'{cwd}\Dokumenty\{word_file}.docx','rb') as f:
					document = Document(f)
				print(f"Document {word_file} found")
			except:
				print(f"Document {word_file} not found")
			
			for section in document.sections:			#loop sekce -> zapati
				footer = section.footer
				for key in stary_veci.keys():
					for foot in footer.paragraphs:	
						self.doc_replace_fuc(dic_new=self.project_dict,dic_old=stary_veci,key=key,place=foot)
						

			for paragraph in document.paragraphs:			#loop paragrafy
				# print (paragraph.text)
				for key in stary_veci.keys():				#loop key z dict
					self.doc_replace_fuc(dic_new=self.project_dict,dic_old=stary_veci,key=key,place=paragraph)

			document.save(rf'C:\Users\{user_name}\Desktop\{word_file}_{self.project_dict["k.u."]}.docx')

		pruvodni_zprava_files = [
			"A_Průvodní zpráva",
			"B_Souhrnná_technická_zpráva",
			"C_Situační_výkresy",
			"D_Dokumentace_objektů_a_technických_a_technologických_zařízení",
			"E_Organizace_výstavby"]		
					
		for file in pruvodni_zprava_files:
			sub_pruvodni_zprava(file)
				
		return	1

	

		


class zadani:
	def __init__(self, project_info):
		# self.cislo_projektu = cislo_projektu
		self.project_f_dir = project_info["project_f_dir"]
		with open("setting.json","r",encoding = 'utf-8') as f:
			file = f.read()
			self.setting = json.loads(file)

	######################## Hledani zadani (file path)######################################
	def hledani_zadani(self):	

		project_f_dir = self.project_f_dir

		try:
			for file in os.listdir(project_f_dir):
				# print (file)
				if 'ZAD' in file:
					zadani = project_f_dir + rf'\{file}'
					self.zadani = zadani
					# print('Zadani:' + zadani)
					return self.zadani
				if 'podklady' == file:
					project_f_dir = project_f_dir +r'\podklady'
					podklady = True
					# print('podklady True')
			
			if os.listdir(project_f_dir) != None and 'podklady' not in locals():
				project_f_dir = project_f_dir +rf'\{file}'
				for file in os.listdir(project_f_dir):
					if 'ZAD' in file:
						zadani = project_f_dir + rf'\{file}'
						self.zadani = zadani
						# print('Zadani:' + zadani)
						return self.zadani
					if 'podklady' in file:
						project_f_dir = project_f_dir +r'\podklady'
						podklady = True
						# print('podklady True')

		except:
			print ('Složka nenalezena')
			return None

		if 'podklady' in locals() and True:
			for file in os.listdir(project_f_dir):
				if 'ZAD' in file:
					zadani = project_f_dir + rf'\{file}'
					self.zadani = zadani
					return self.zadani
					# print('Zadani:' + zadani)
				elif 'ZN' == file:
					project_f_dir = project_f_dir +r'\ZN'
					ZN = True
					# print('ZN True')
		
		if 'ZN' in locals() and True:
		# print(project_f_dir)
			for file in os.listdir(project_f_dir):
				# print(file)
				if 'ZAD' in file:
					zadani = project_f_dir + rf'\{file}'
					self.zadani = zadani
					return self.zadani
					print('Zadani:' + zadani)

			if os.listdir(project_f_dir) != None:
				# print ('Not none')
				for file in os.listdir(project_f_dir):	
					find_file_directory = project_f_dir +rf'\{file}'
					
					try:
						for f in os.listdir(find_file_directory):
							
							if 'ZAD' in f:
								zadani = find_file_directory + rf'\{f}'
								self.zadani = zadani
								return self.zadani
					except:	#pass all not file
						pass
		
		
	######################### zadani info #########################################
	def info_zadani(self):
		zadani_file = self.hledani_zadani()
				

		if zadani_file is None:
			print ('File None')
			return None

		with open(zadani_file,'rb') as f:
			document = Document(f)

		dict_extract = self.setting["zadani"]

		dict_ZAD = {}

		for index,paragraph in enumerate(document.paragraphs):			#loop paragrafy
			
			for key in dict_extract.keys():				#loop key z dict
				# print(key, paragraph.text)
				if str(dict_extract.get(f"{key}")) in paragraph.text:			# key value v text paragrafu
					# print(index, key,str(dict_extract.get(f"{key}")) in paragraph.text)
					inline = paragraph.runs									#vestavena funkce loop z modulu doxc pro udryeni stylu
					# print(index, 'inline:', inline)

					for text in inline:						#text
						if str(dict_extract.get(f"{key}")) in text.text and key == 'termin_realizace':
							for n in inline:
								if any(x.isdigit() for x in n.text) :
									termin = n.text
									ctvrleti = (int(termin.split('.')[1])-1)//3+1
									termin = str(ctvrleti) + '.čtvrletí ' + n.text.split('.')[2]
									dict_ZAD.setdefault(key, termin)
									# print (str(dict_extract.get(f"{key}")), n.text)

					if key == 'cena_i':
						pocatek = index + 1

					if key == 'cena_f':
						konec = index
		


		
		for n in range(pocatek, konec):
			para = document.paragraphs[n]

			if  any(x.isdigit() for x in para.text) :
				dict_ZAD.setdefault("cena", para.text + ' tis.Kč')
			# print(para.text)
		

		self.dict_ZAD = dict_ZAD
		return self.dict_ZAD

	



######################### search 2 json #######################################
def search_k_u_ID_2json():
	list_url = ['https://www.cuzk.cz/Dokument.aspx?AKCE=META:SESTAVA:MDR001_XSLT:WEBCUZK_KRAJEKOD:560_PZEROK:2022',
	'https://www.cuzk.cz/Dokument.aspx?AKCE=META:SESTAVA:MOR001_XSLT:WEBCUZK_KRAJE:560_PZER:2022',
	'https://www.cuzk.cz/Dokument.aspx?AKCE=META:SESTAVA:MDR001_XSLT:WEBCUZK_KRAJEKOD:200_PZEROK:2022',
	'https://www.cuzk.cz/Dokument.aspx?AKCE=META:SESTAVA:MOR001_XSLT:WEBCUZK_KRAJE:200_PZER:2022'
	]

	dict_ku={}
	
	for URL in list_url:
		page = requests.get(URL)
		soup = BeautifulSoup(page.content, "html.parser")
		results = soup.find("tbody")
		for div in results.find_all("td", class_="obec"):
			nazev_k_u = div.find('a').contents[0]
			ID_k_u = div.find('a')['href'].split(':')[-1]
			dict_ku.setdefault(nazev_k_u, ID_k_u)

	json_con = json.dumps(dict_ku,indent=4, sort_keys=True)
	
	with open('k_u_ID.json', 'w') as f:
		f.write(json_con)		

######################################################################	 



###################### Hledani URL function #######################
	# def hledani_URL(informace_o_projektu):	#vlozit dict z search

	# 	list_URL = ("https://www.cuzk.cz/Dokument.aspx?AKCE=META:SESTAVA:MDR001_XSLT:WEBCUZK_KRAJEKOD:560_PZEROK:2022",
	# 		"https://www.cuzk.cz/Dokument.aspx?AKCE=META:SESTAVA:MDR001_XSLT:WEBCUZK_KRAJEKOD:200_PZEROK:2022"
	# 			)

	# 	informace_jmena = informace_o_projektu.get("nazev_projektu")	#parsel nazev mesta
	# 	informace_jmena = informace_jmena.replace(" - ","-")	#odstraneni mezery
	# 	informace_jmena = re.split('[-_,]',informace_jmena)		#split 
	# 	print(informace_jmena, informace_jmena[1])

	# 	if len(informace_jmena[0]) > 2:
	# 		informace_jmena.insert(0," ")
	# 		nazev_mesta = informace_jmena[1]
	# 	else:
	# 		nazev_mesta = informace_jmena[1]

		

	# 	if "Ústí n.L" in nazev_mesta:
	# 		nazev_mesta = nazev_mesta.replace("Ústí n.L.","Ústí nad Labem")


	# 	for URL in list_URL:
	# 		page = requests.get(URL)
	# 		soup = BeautifulSoup(page.content, "html.parser")
	# 		results = soup.find("tbody")
			
	# 		for div in results.find_all("td", class_="obec"):

	# 			if nazev_mesta == div.find('a').contents[0]:
	# 				link_k_u = 'https://www.cuzk.cz/'+div.find('a')['href'] #hledano link na k.u.
	# 				break

	# 	try:
	# 		return link_k_u, nazev_mesta # URL
	# 	except UnboundLocalError:
	# 		print ("k.u. Nenalezeno")
	# 		return None



# def pdf_cover_content():
# 	o = client.Dispatch("Excel.Application")
# 	o.Visible = False
# 	wb_path = rf'C:\Users\David\Desktop\New Folder\py test\test file.xlsx'

# 	wb = o.Workbooks.Open(wb_path)


# 	#say you want to print these sheets
# 	ws_index_list_cover_pages = [1] 
	



# def pdf_cover_merger(a,index,name=None,b='BlankPDF.pdf',):
# 	pdfs = [a, b]

# 	merger = PdfFileMerger()

# 	for pdf in pdfs:
# 		merger.append(pdf)

# 	merger.write(f"{name}{index}.pdf")
# 	merger.close()
# 	os.remove(f'{a}')	

plna_moc = {
	"plna_moc_CEZ" : r'C:\Users\David\OneDrive - R-built s.r.o\Dokumenty\1_plné moci\PLNÁ MOC EV ČÍSLO PM - 031_2021.pdf',
	"plna_moc_LS" :  r'C:\Users\David\OneDrive - R-built s.r.o\Dokumenty\1_plné moci\Plná moc LS.pdf',
	"plna_moc_KS" : r'C:\Users\David\OneDrive - R-built s.r.o\Dokumenty\1_plné moci\Plná moc KS.pdf',
	"CEZ_Sou_Stan" : r'C:\Users\David\Desktop\New Folder\Work thing\ČEZ souhlasné stanovisko - PRO VŠECHNY.pdf'

}
	
###################        Nejak nefunguje       #######################
# def pdf_parametr(a):
# 		ws = wb.Worksheets[index - 1]
# 		ws.PageSetup.Zoom = False
# 		ws.PageSetup.FitToPagesTall = 1
# 		ws.PageSetup.FitToPagesWide = 1
# 		ws.PageSetup.PrintArea = a

# 	for index in ws_index_list_content_pages:
# 		#content list arep
# 		pdf_parametr(a='A1:D20')

# 	#export cover pages	
# 	for n in ws_index_list_content_pages:
# 		wb.WorkSheets(n).Select()
# 		wb.ActiveSheet.ExportAsFixedFormat(0, rf'C:\Users\David\Desktop\New Folder\py test\obsah{n-1}_temp.pdf')
# 		pdf_cover_merger(a=f'obsah{n-1}_temp.pdf',index=n-1,name='obsah_t')
	
# 	for index in ws_index_list_cover_pages:
# 		#cover list are
# 		pdf_parametr(a='A1:J50')

# 	#export content pages
# 	for n in ws_index_list_cover_pages:
# 		wb.WorkSheets(n).Select()
# 		wb.ActiveSheet.ExportAsFixedFormat(0, rf'C:\Users\David\Desktop\New Folder\py test\doklad_cover{n}.pdf')
# 		pdf_cover_merger(a=f'doklad_cover{n}.pdf',index=n,name='doklad_t')
	
# 	#merge pages
# 	for n in ws_index_list_cover_pages:
# 		pdfs = [f'doklad_t{n}.pdf', f'obsah_t{n}.pdf']

# 		merger = PdfFileMerger()

# 		for pdf in pdfs:
# 			merger.append(pdf)

# 		merger.write(f"doklad{n}.pdf")
# 		merger.close()
# 		os.remove(f'doklad_t{n}.pdf')
# 		os.remove(f'obsah_t{n}.pdf')



if __name__ == '__main__':
	test()